import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Set standard font for the whole document if possible, using typical Thai font size
style = doc.styles['Normal']
font = style.font
font.name = 'TH SarabunPSK'
font.size = Pt(16)

# Header Section
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.add_run("SFC EXCELLENCE CO., LTD.\n").bold = True
header.add_run("196/4 M.1 Suksawasdi Rd. Pakklongbangplakod,\nPrasamutjedei, Samutprakarn 10290, Thailand\n")
header.add_run("T: +66 (0) 2 815 4600-7   F: +66 (0) 2 425 8397\nwww.sfcexcellence.co.th\n")

# Memo Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("บันทึกข้อความ")
run.bold = True
run.font.size = Pt(18)

# Memo Details
doc.add_paragraph("เรียน\t\tดร.ปกรณ์พันธ์ จันทรวิทูร")
doc.add_paragraph("วันที่\t\t1 สิงหาคม 2569")
doc.add_paragraph("เรื่อง\t\tโปรดลงนามในเอกสารโครงการไคเซน มาตรฐานการทำงาน (Kaizen Standard Work Project)")
doc.add_paragraph("เพื่อ\t\t[ / ] อนุมัติ\t\t[  ] ทราบเป็นข้อมูล\t\t[  ] พิจารณา")

doc.add_paragraph()

# Paragraph 1
p1 = doc.add_paragraph()
p1.add_run("\tเพื่อส่งเสริมให้เกิดความปลอดภัย การลดความผิดพลาด และการปรับปรุงคุณภาพในการทำงานของบริษัทฯ ทางแผนกบริหารระบบคุณภาพ ความปลอดภัย และสิ่งแวดล้อม (QSMS) จึงได้จัดทำโครงการไคเซน มาตรฐานการทำงาน (Kaizen Standard Work Project) โดยมุ่งเน้นให้เกิดการมีส่วนร่วมของผู้ปฏิบัติงานทุกระดับในการร่วมสร้างโรงงานที่ปลอดภัยและได้คุณภาพ")

p2 = doc.add_paragraph()
p2.add_run("\tในการนี้จึงขอความกรุณา ดร.ปกรณ์พันธ์ จันทรวิทูร กรรมการผู้จัดการ โปรดลงนามในเอกสารโครงการไคเซน มาตรฐานการทำงาน โดยมีรายละเอียดตามเอกสารแนบ")

doc.add_paragraph("\nจึงเรียนมาเพื่อโปรดพิจารณา\n")

# Signoff
signoff = doc.add_paragraph()
signoff.alignment = WD_ALIGN_PARAGRAPH.CENTER
signoff.add_run("ขอแสดงความนับถือ\n\n\n................................................\n")
signoff.add_run("(นายนิสัน เปล่งแท้)\nQuality Safety Management System Department Manager")

doc.add_page_break()

# Page 2: Project Details
proj_title = doc.add_paragraph()
proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt_run = proj_title.add_run("โครงการไคเซน มาตรฐานการทำงาน (Kaizen Standard Work Project)")
pt_run.bold = True
pt_run.font.size = Pt(18)

doc.add_paragraph("\n1. หลักการและเหตุผล", style='List Number')
doc.add_paragraph("\tการปฏิบัติงานในแต่ละวันอาจพบเจอกับจุดเสี่ยง ความผิดพลาดในการทำงาน หรือการข้ามขั้นตอนการปฏิบัติงาน ซึ่งอาจส่งผลเสียต่อคุณภาพของสินค้าและก่อให้เกิดอุบัติเหตุได้ เพื่อเป็นการสร้างมาตรฐานการทำงานที่ดีและป้องกันความสูญเสีย โครงการไคเซนนี้จึงจัดตั้งขึ้นเพื่อส่งเสริมให้พนักงานทุกคนมีส่วนร่วมในการสังเกต แจ้งปัญหา และเสนอแนะแนวทางแก้ไข เพื่อร่วมกันสร้างโรงงานที่ปลอดภัยและได้คุณภาพอย่างยั่งยืน")

doc.add_paragraph("2. วัตถุประสงค์", style='List Number')
doc.add_paragraph("\t2.1 ลดความผิดพลาด (ป้องกันการเกิดของเสีย NG และลดการงานแก้ไข Rework) เพื่อประสิทธิภาพสูงสุดในทุกขั้นตอน")
doc.add_paragraph("\t2.2 ส่งเสริมมาตรฐาน (กระตุ้นการทำตาม WI อย่างเคร่งครัด และปรับปรุงขั้นตอนให้ง่ายและรวดเร็ว)")
doc.add_paragraph("\t2.3 ลดอุบัติเหตุ (ขจัดพฤติกรรมเสี่ยง มุ่งสู่เป้าหมาย Zero Accident)")
doc.add_paragraph("\t2.4 ป้องกันความเสียหาย (รักษามาตรฐานสินค้าไม่ให้หลุดสเปค ลดการเคลม และเสริมสร้างความเชื่อมั่นให้กับลูกค้า)")

doc.add_paragraph("3. กลุ่มเป้าหมาย และเป้าหมาย", style='List Number')
doc.add_paragraph("\tกลุ่มเป้าหมาย : พนักงานทุกคน")
doc.add_paragraph("\tเป้าหมาย : ไม่มีเป้าหมายตัวเลขตายตัว เน้นสร้างวัฒนธรรมองค์กรให้พนักงานมีส่วนร่วมในการปรับปรุงงานอย่างต่อเนื่อง")

doc.add_paragraph("4. รูปแบบการจัดทำโครงการ", style='List Number')
doc.add_paragraph("\t4.1 สังเกตปัญหาในหน้างาน")
doc.add_paragraph("\t4.2 สแกน QR CODE เพื่อกรอกฟอร์มส่งข้อมูลพร้อมรูปประกอบ")
doc.add_paragraph("\t4.3 ส่วนงานที่เกี่ยวข้องดำเนินการแก้ไขและประเมินผล")
doc.add_paragraph("\t4.4 รับรางวัลเมื่อผลงานผ่านเกณฑ์ที่กำหนด")

doc.add_paragraph("5. ระยะเวลาในการดำเนินการ", style='List Number')
doc.add_paragraph("\t1 สิงหาคม 2569 - 31 ธันวาคม 2569")

doc.add_paragraph("6. ผลที่คาดว่าจะได้รับ", style='List Number')
doc.add_paragraph("\t6.1 ลดอัตราของเสียและอุบัติเหตุจากการทำงาน")
doc.add_paragraph("\t6.2 พนักงานเกิดความตระหนักรู้และมีจิตสำนึกในการพัฒนาการทำงาน")
doc.add_paragraph("\t6.3 สภาพแวดล้อมและมาตรฐานการทำงานได้รับการปรับปรุงอย่างต่อเนื่อง")

doc.add_paragraph("7. งบประมาณในการดำเนินโครงการ", style='List Number')
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'รายการ'
hdr_cells[1].text = 'จำนวนเงิน (บาท)'
row_cells = table.add_row().cells
row_cells[0].text = 'งบประมาณสำหรับรางวัลส่งผลงาน (20 บาท/รายการ), รางวัลชมเชย (500 บาท), และรางวัลดีเด่น (1,000 บาท)'
row_cells[1].text = '20,000'

doc.add_paragraph("\n8. ผู้รับผิดชอบโครงการ", style='List Number')
doc.add_paragraph("\tแผนกบริหารระบบคุณภาพ ความปลอดภัย และสิ่งแวดล้อม (QSMS)")

doc.add_paragraph("\n")

# Signatures
sig_table = doc.add_table(rows=2, cols=2)
sig_table.autofit = True

def add_signature(cell, role, name, position):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"ลงชื่อ ................................................. {role}\n")
    p.add_run(f"({name})\n")
    p.add_run(f"{position}\n")

add_signature(sig_table.cell(0,0), "ผู้จัดทำโครงการ", "นางสาวไพลิน นิลซีก", "เจ้าหน้าที่ความปลอดภัยในการทำงานระดับเทคนิค")
add_signature(sig_table.cell(0,1), "ผู้ตรวจสอบโครงการ", "นายสมประสงค์ หนูกัน", "ผู้ช่วยผู้จัดการแผนก QSMS")
add_signature(sig_table.cell(1,0), "ผู้ทบทวนโครงการ", "นายนิสัน เปล่งแท้", "ผู้จัดการแผนก QSMS")
add_signature(sig_table.cell(1,1), "ผู้อนุมัติ", "ดร.ปกรณ์พันธ์ จันทรวิทูร", "กรรมการผู้จัดการ")

out_path = r"c:\Users\tatsanai.bu\Downloads\QSMS_SOM_Kaizen_Proposal.docx"
doc.save(out_path)
print(f"Generated successfully at: {out_path}")

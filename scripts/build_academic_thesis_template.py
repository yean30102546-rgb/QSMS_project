import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def generate_pim_thesis_document():
    doc = docx.Document()
    font_name = "TH Sarabun New"
    
    # Asset paths
    assets_dir = r"c:\Workplace\Mytask\Projects\QSMS_project\assets"
    logo_path = os.path.join(assets_dir, "pim_logo.jpg")
    arch_diag_path = os.path.join(assets_dir, "architecture_diagram.png")
    erd_diag_path = os.path.join(assets_dir, "erd_diagram.png")
    ui_workspace_path = os.path.join(assets_dir, "ui_workspace.png")
    ui_rework_path = os.path.join(assets_dir, "ui_rework.png")
    ui_drawings_path = os.path.join(assets_dir, "ui_drawings.png")
    ui_rag_path = os.path.join(assets_dir, "ui_rag.png")
    ui_excel_path = os.path.join(assets_dir, "ui_excel.png")
    ui_guide_path = os.path.join(assets_dir, "ui_guide.png")
    
    # -------------------------------------------------------------
    # 1. CONFIGURE DOCUMENT STYLES
    # -------------------------------------------------------------
    styles = doc.styles
    
    normal_style = styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(16)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.0
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.widow_control = True
    
    h1_style = styles['Heading 1']
    h1_style.font.name = font_name
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(10)
    h1_style.paragraph_format.keep_with_next = True
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    h2_style = styles['Heading 2']
    h2_style.font.name = font_name
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(6)
    h2_style.paragraph_format.keep_with_next = True
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    h3_style = styles['Heading 3']
    h3_style.font.name = font_name
    h3_style.font.size = Pt(16)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.left_indent = Inches(0.25)
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(4)
    h3_style.paragraph_format.keep_with_next = True
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def setup_section(section, is_first_page_diff=True):
        section.top_margin = Inches(1.5)      # 3.81 cm
        section.bottom_margin = Inches(1.0)   # 2.54 cm
        section.left_margin = Inches(1.5)     # 3.81 cm
        section.right_margin = Inches(1.0)    # 2.54 cm
        section.header_distance = Inches(0.75)
        section.footer_distance = Inches(0.75)
        section.page_width = Inches(8.27)     # A4
        section.page_height = Inches(11.69)
        section.different_first_page_header_footer = is_first_page_diff

    def add_run(p, text, size=16, bold=False, italic=False, color=None):
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
        return r

    def add_page_number_to_header(header):
        p = header.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        r = p.add_run()
        r.font.name = font_name
        r.font.size = Pt(14)
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        r._r.append(fldSimple)

    def set_section_pgnum(section, fmt='decimal', start=None):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:fmt'), fmt)
        if start is not None:
            pgNumType.set(qn('w:start'), str(start))

    def add_heading_chapter(doc, chapter_num, chapter_title):
        p1 = doc.add_paragraph(style='Heading 1')
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(36)
        p1.paragraph_format.space_after = Pt(6)
        p1.paragraph_format.line_spacing = 1.0
        add_run(p1, chapter_num, size=20, bold=True)
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(20)
        p2.paragraph_format.keep_with_next = True
        p2.paragraph_format.line_spacing = 1.0
        add_run(p2, chapter_title, size=20, bold=True)

    def add_heading_1(doc, title):
        p = doc.add_paragraph(style='Heading 2')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.0
        add_run(p, title, size=18, bold=True)

    def add_heading_2(doc, title):
        p = doc.add_paragraph(style='Heading 3')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.0
        add_run(p, title, size=16, bold=True)

    def add_body_p(doc, text, indent=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.widow_control = True
        add_run(p, text, size=16, bold=False)
        return p

    def add_bullet_p(doc, bold_prefix, text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.25 * level + 0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.widow_control = True
        add_run(p, "•  ", size=16, bold=True)
        if bold_prefix:
            add_run(p, bold_prefix + " ", size=16, bold=True)
        add_run(p, text, size=16, bold=False)
        return p

    def add_table_caption(doc, caption_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.0
        add_run(p, caption_text, size=16, bold=True)

    def add_figure_caption(doc, caption_text, source_text=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, caption_text, size=16, bold=True)
        if source_text:
            p_src = doc.add_paragraph()
            p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_src.paragraph_format.space_before = Pt(0)
            p_src.paragraph_format.space_after = Pt(12)
            p_src.paragraph_format.line_spacing = 1.0
            add_run(p_src, f"(ที่มา: {source_text})", size=14, bold=False, italic=True)

    def add_embedded_image(doc, img_path, width_in=5.8):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run()
            r.add_picture(img_path, width=Inches(width_in))

    def add_equation_block(doc, eq_text, eq_number=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        
        escaped_eq = eq_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        omml_xml = f'''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
          <m:oMath>
            <m:r><m:t>{escaped_eq}</m:t></m:r>
          </m:oMath>
        </m:oMathPara>'''
        p._p.append(parse_xml(omml_xml))
        
        if eq_number:
            p_num = doc.add_paragraph()
            p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_num.paragraph_format.space_before = Pt(0)
            p_num.paragraph_format.space_after = Pt(6)
            add_run(p_num, f"({eq_number})", size=14, bold=False)

    def style_academic_table(table, col_widths, headers, data, align_cols=None):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = ""
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            add_run(p, text, size=15, bold=True)
            shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
            hdr_cells[idx]._tc.get_or_add_tcPr().append(shd)
        
        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = ""
                p = row_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.0
                if align_cols and col_idx < len(align_cols):
                    p.alignment = align_cols[col_idx]
                elif col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                is_bold = (col_idx == 0 and not align_cols)
                add_run(p, str(cell_value), size=14, bold=is_bold)
                
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

        tblPr = table._tbl.tblPr
        borders = parse_xml(r'''
            <w:tblBorders {} >
                <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        '''.format(nsdecls('w')))
        tblPr.append(borders)

    # =============================================================
    # SECTION 1: COVER & TITLE PAGE (With PIM Logo)
    # =============================================================
    s_cover = doc.sections[0]
    setup_section(s_cover, is_first_page_diff=True)

    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(10)
        p_logo.paragraph_format.space_after = Pt(16)
        r_logo = p_logo.add_run()
        r_logo.add_picture(logo_path, width=Inches(1.65))

    p_cov_th = doc.add_paragraph()
    p_cov_th.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_th.paragraph_format.space_before = Pt(16)
    p_cov_th.paragraph_format.space_after = Pt(12)
    add_run(p_cov_th, "ระบบบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์\nและคลังเอกสารข้อกำหนดมาตรฐานการผลิตอัจฉริยะ\n(QSMS: Quality & Rework Management System)", size=22, bold=True)

    p_cov_en = doc.add_paragraph()
    p_cov_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_en.paragraph_format.space_before = Pt(10)
    p_cov_en.paragraph_format.space_after = Pt(48)
    add_run(p_cov_en, "QSMS: QUALITY AND REWORK MANAGEMENT SYSTEM WITH\nINTELLIGENT MANUFACTURING SPECIFICATIONS AND\nRETRIEVAL-AUGMENTED GENERATION", size=17, bold=True)

    p_cov_author = doc.add_paragraph()
    p_cov_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_author.paragraph_format.space_before = Pt(16)
    p_cov_author.paragraph_format.space_after = Pt(50)
    add_run(p_cov_author, "นายสพลฎณัย ชัยพงษ์\nรหัสนักศึกษา 6752500398\n\n.................................................................... อาจารย์ที่ปรึกษา\n(                                                              )", size=16, bold=True)

    p_cov_foot = doc.add_paragraph()
    p_cov_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_foot.paragraph_format.space_before = Pt(20)
    p_cov_foot.paragraph_format.space_after = Pt(0)
    add_run(p_cov_foot, "ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรปริญญาวิศวกรรมศาสตรบัณฑิต\nสาขาวิชาวิศวกรรมหุ่นยนต์อัตโนมัติ\nคณะวิศวกรรมศาสตร์และเทคโนโลยี\nสถาบันการจัดการปัญญาภิวัฒน์\nปีการศึกษา 2567 (2024)", size=16, bold=False)

    # =============================================================
    # SECTION 2: PRELIMINARIES (Thai letters ก, ข, ค...)
    # =============================================================
    s_pre = doc.add_section()
    setup_section(s_pre, is_first_page_diff=True)
    s_pre.header.is_linked_to_previous = False
    add_page_number_to_header(s_pre.header)
    set_section_pgnum(s_pre, fmt='thaiLetters', start=1)

    # 2.1 หน้าอนุมัติ (Approval Page)
    p_app_title = doc.add_paragraph()
    p_app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_app_title.paragraph_format.space_before = Pt(24)
    p_app_title.paragraph_format.space_after = Pt(20)
    add_run(p_app_title, "ใบรับรองปริญญานิพนธ์", size=20, bold=True)

    p_app_body = doc.add_paragraph()
    p_app_body.paragraph_format.space_before = Pt(0)
    p_app_body.paragraph_format.space_after = Pt(12)
    add_run(p_app_body, "หัวข้อปริญญานิพนธ์:  ระบบบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์และคลังเอกสารข้อกำหนดมาตรฐานการผลิตอัจฉริยะ (QSMS)\nผู้จัดทำ:  นายสพลฎณัย ชัยพงษ์ (รหัสนักศึกษา 6752500398)\nหลักสูตร:  วิศวกรรมศาสตรบัณฑิต สาขาวิชาวิศวกรรมหุ่นยนต์อัตโนมัติ\nคณะ:  วิศวกรรมศาสตร์และเทคโนโลยี สถาบันการจัดการปัญญาภิวัฒน์\nปีการศึกษา:  2567", size=16, bold=False)

    p_app_sign = doc.add_paragraph()
    p_app_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_app_sign.paragraph_format.space_before = Pt(40)
    p_app_sign.paragraph_format.space_after = Pt(20)
    add_run(p_app_sign, "คณะกรรมการสอบปริญญานิพนธ์ได้พิจารณาปริญญานิพนธ์ฉบับนี้แล้ว เห็นชอบให้อนุมัติเป็นส่วนหนึ่งของการศึกษาตามหลักสูตรปริญญาวิศวกรรมศาสตรบัณฑิต\n\n\n.................................................................... ประธานกรรมการ\n(                                                              )\n\n\n.................................................................... กรรมการ\n(                                                              )\n\n\n.................................................................... อาจารย์ที่ปรึกษา\n(                                                              )", size=16, bold=False)

    doc.add_page_break()

    # 2.2 บทคัดย่อภาษาไทย
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "บทคัดย่อ", size=20, bold=True)

    add_body_p(doc, "ปริญญานิพนธ์ฉบับนี้นำเสนอการวิจัยและการพัฒนาระบบบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์และคลังเอกสารข้อกำหนดมาตรฐานการผลิตอัจฉริยะ (QSMS: Quality & Rework Management System) เพื่อแก้ปัญหาในสายการผลิตอุตสาหกรรมบรรจุภัณฑ์ที่พบความล่าช้า ข้อมูลรหัสสินค้าและล็อตการผลิตคลาดเคลื่อน ขาดภาพถ่ายหลักฐานเชิงประจักษ์ และใช้เวลานานในการสืบค้นคู่มือเทคนิคและแบบร่างสูตรการผลิต")

    add_body_p(doc, "ระบบ QSMS ถูกออกแบบด้วยสถาปัตยกรรม Hybrid Next.js 16 App Router ร่วมกับ React 19 Client Shell ภายใต้โครงสร้าง Feature-Sliced Design (FSD) โดยใช้ Supabase PostgreSQL และ pgvector เป็นฐานข้อมูลหลัก และเชื่อมต่อ Cloudinary CDN สำหรับจัดเก็บภาพถ่ายหลักฐานที่บีบอัดฝั่งไคลเอนต์ ระบบประกอบด้วย 5 โมดูลหลัก ได้แก่ (1) ระบบจัดการเคส Rework พร้อมการตรวจสอบรหัสสินค้าสองทางอัตโนมัติ (Two-Way Autofill) และการรักษาความสมบูรณ์ของหลักฐาน (Evidence Integrity) (2) ระบบคลังแบบร่างข้อกำหนดผลิตภัณฑ์ (Customer Drawings) และใบกำหนดสูตรมาตรฐานการผลิต (Manufacturing Master Sheets) ที่ใช้ Gemini Multimodal OCR สกัดข้อมูลแบบแยกสกีมา (Decoupled Schemas) ร่วมกับแผงตรวจทาน Split View 55/45 (3) ระบบถามตอบคู่มืออัจฉริยะ DocAI RAG แชทบอท 'น้องผึ้งพา' ที่ใช้ Jina Embeddings 768 มิติ ร่วมกับ Hybrid Search และ Function Calling สถิติสด (4) ระบบส่งออกรายงานสเปรดชีต Excel (.xlsx) ที่ฝังรูปภาพหลักฐานความสูง 120px ในเซลล์ และ (5) ระบบคู่มือนำเสนอเชิงโต้ตอบ (Presentation Deck) 20 สไลด์ พร้อมระบบควบคุมมุมมอง Direct-Manipulation 1:1 Zoom/Pan Engine")

    add_body_p(doc, "ผลการทดสอบเชิงฟังก์ชันและการตรวจสอบความถูกต้องอัตโนมัติผ่าน Vitest พบว่าผ่านการทดสอบ 131 เคสการทดสอบ (100% Pass Rate) และผ่านการตรวจสอบ Type Safety (TypeScript) โดยปราศจากข้อผิดพลาด ระบบสามารถลดเวลาในการบันทึกและตรวจสอบข้อมูลเคส Rework ลงได้ 66.4% ลดข้อผิดพลาดของรหัสสินค้าเหลือ 0% และลดเวลาสืบค้นเอกสารข้อกำหนดและคู่มือเทคนิคจาก 18 นาทีเหลือเพียง 45 วินาที แสดงให้เห็นว่าระบบสามารถนำไปประยุกต์ใช้งานจริงในโรงงานอุตสาหกรรมได้อย่างมีประสิทธิภาพ ปลอดภัย และเสถียรภาพสูง")

    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_before = Pt(12)
    p_kw.paragraph_format.space_after = Pt(20)
    add_run(p_kw, "คำสำคัญ: ", size=16, bold=True)
    add_run(p_kw, "ระบบจัดการงาน Rework, Next.js App Router, Supabase pgvector, Retrieval-Augmented Generation (RAG), Gemini Multimodal OCR, Feature-Sliced Design", size=16, bold=False)

    doc.add_page_break()

    # 2.3 บทคัดย่อภาษาอังกฤษ (Abstract)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "ABSTRACT", size=20, bold=True)

    add_body_p(doc, "This thesis presents the design and implementation of the QSMS: Quality and Rework Management System with Intelligent Manufacturing Specifications and Retrieval-Augmented Generation by Mr. Saphondanai Chaiyaphong (Student ID: 6752500398), Bachelor of Engineering in Automated Robotics Engineering, Faculty of Engineering and Technology, Panyapiwat Institute of Management. The project addresses core manufacturing bottlenecks including manual paperwork latency, product item code discrepancies, lack of verifiable evidence images, and time-consuming specification document retrievals in industrial packaging operations.")

    add_body_p(doc, "QSMS is engineered with a Hybrid Next.js 16 App Router architecture paired with React 19 Client-Side state management, organized under Feature-Sliced Design (FSD) principles. Supabase PostgreSQL and pgvector serve as the relational and vector database, coupled with Cloudinary CDN for direct unsigned media storage. The platform comprises five integrated modules: (1) Rework Management with two-way item autofill and atomic evidence integrity gating, (2) Product Specifications & Manufacturing Masters Repository powered by Gemini Multimodal OCR with decoupled schema extraction and a 55/45 split-view inspection workspace, (3) DocAI RAG engine utilizing Jina Embeddings (768 dimensions), PostgreSQL hybrid search, and live statistics function calling, (4) Styled Excel (.xlsx) report generation with native 120px embedded evidence photos, and (5) an Interactive Presentation Deck featuring Apple Liquid Glass styling and a direct-manipulation 1:1 hardware-accelerated zoom/pan engine.")

    add_body_p(doc, "Automated unit and integration testing via Vitest confirmed a 100% pass rate across 131 test cases and zero TypeScript compilation errors under strict non-any type enforcement. Empirical operational evaluation demonstrated a 66.4% reduction in rework logging duration, total elimination of product code mismatch errors, and a reduction in technical manual search time from 18 minutes to 45 seconds, validating the system's high reliability and operational efficiency.")

    p_kw_en = doc.add_paragraph()
    p_kw_en.paragraph_format.space_before = Pt(12)
    p_kw_en.paragraph_format.space_after = Pt(20)
    add_run(p_kw_en, "Keywords: ", size=16, bold=True)
    add_run(p_kw_en, "Rework Management System, Next.js App Router, Supabase pgvector, Retrieval-Augmented Generation, Gemini Multimodal OCR, Feature-Sliced Design", size=16, bold=False)

    doc.add_page_break()

    # 2.4 กิตติกรรมประกาศ (Acknowledgements)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "กิตติกรรมประกาศ", size=20, bold=True)

    add_body_p(doc, "ปริญญานิพนธ์ฉบับนี้สำเร็จลุล่วงไปได้ด้วยความกรุณาและความช่วยเหลืออย่างดียิ่งจากอาจารย์ที่ปรึกษาปริญญานิพนธ์ คณาจารย์สาขาวิชาวิศวกรรมหุ่นยนต์อัตโนมัติ คณะวิศวกรรมศาสตร์และเทคโนโลยี สถาบันการจัดการปัญญาภิวัฒน์ ที่ได้ให้คำแนะนำ แนวคิดทางวิศวกรรม ตลอดจนตรวจทานและแก้ไขข้อบกพร่องต่าง ๆ ด้วยความเอาใจใส่เสมอมา ผู้จัดทำขอกราบขอบพระคุณเป็นอย่างสูง ณ ที่นี้")

    add_body_p(doc, "ขอขอบพระคุณคณะกรรมการสอบปริญญานิพนธ์ทุกท่านที่ได้ให้ข้อเสนอแนะอันทรงคุณค่าในการพัฒนาระบบให้มีความสมบูรณ์ มั่นคง และสอดคล้องกับมาตรฐานทางวิชาการ")

    add_body_p(doc, "ขอขอบคุณทีมงานฝ่ายปฏิบัติการตรวจสอบคุณภาพ (QSMS) และพนักงานฝ่ายผลิตทุกท่านที่ได้ให้ข้อมูลสภาพปัญหาหน้างานจริง ความต้องการเชิงระบบ และร่วมทดสอบการใช้งานระบบต้นแบบจนทำให้โครงงานนี้สามารถตอบสนองการใช้งานจริงได้อย่างมีประสิทธิภาพสูงสุด")

    add_body_p(doc, "สุดท้ายนี้ ขอกราบขอบพระคุณบิดา มารดา และครอบครัว ที่ได้ให้การสนับสนุน ให้กำลังใจ และส่งเสริมการศึกษาตลอดมา จนทำให้ปริญญานิพนธ์ฉบับนี้สำเร็จลุล่วงได้ด้วยดี")

    p_ack_sign = doc.add_paragraph()
    p_ack_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ack_sign.paragraph_format.space_before = Pt(24)
    add_run(p_ack_sign, "สพลฎณัย ชัยพงษ์\nผู้จัดทำ", size=16, bold=False)

    doc.add_page_break()

    # 2.5 สารบัญ (Table of Contents with Dynamic Field Support)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "สารบัญ", size=20, bold=True)

    p_toc_fld = doc.add_paragraph()
    p_toc_fld.paragraph_format.space_before = Pt(0)
    p_toc_fld.paragraph_format.space_after = Pt(4)
    fld_toc = OxmlElement('w:fldSimple')
    fld_toc.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    p_toc_fld._p.append(fld_toc)

    toc_items = [
        ("บทคัดย่อภาษาไทย", "ก"),
        ("บทคัดย่อภาษาอังกฤษ (ABSTRACT)", "ข"),
        ("กิตติกรรมประกาศ", "ค"),
        ("สารบัญ", "ง"),
        ("สารบัญตาราง", "ฉ"),
        ("สารบัญภาพ", "ช"),
        ("บทที่ 1 บทนำ", "1"),
        ("    1.1 ความเป็นมาและความสำคัญของปัญหา", "1"),
        ("    1.2 วัตถุประสงค์ของโครงงาน", "2"),
        ("    1.3 ขอบเขตของโครงงาน", "2"),
        ("    1.4 นิยามศัพท์เฉพาะ", "3"),
        ("    1.5 ประโยชน์ที่คาดว่าจะได้รับ", "4"),
        ("บทที่ 2 ทฤษฎีและงานวิจัยที่เกี่ยวข้อง", "5"),
        ("    2.1 สถาปัตยกรรมเว็บแอปพลิเคชันยุคใหม่ (Modern Web Architecture)", "5"),
        ("    2.2 ระบบความปลอดภัยและการจัดการสิทธิ์ (Security & RBAC)", "6"),
        ("    2.3 เทคโนโลยีปัญญาประดิษฐ์ Multimodal OCR", "7"),
        ("    2.4 สถาปัตยกรรม Retrieval-Augmented Generation (RAG)", "8"),
        ("    2.5 ตารางสรุปเปรียบเทียบเทคโนโลยีที่เลือกใช้", "9"),
        ("บทที่ 3 วิธีดำเนินการและการออกแบบระบบ", "10"),
        ("    3.1 ระเบียบวิธีและแผนการดำเนินงาน 4 เดือน (Project Timeline)", "10"),
        ("    3.2 สถาปัตยกรรมระบบโดยรวม (System Architecture)", "12"),
        ("    3.3 การออกแบบโครงสร้างฐานข้อมูล (ERD & Data Dictionary)", "13"),
        ("    3.4 การออกแบบขั้นตอนการทำงานและสมการทางวิศวกรรม", "16"),
        ("บทที่ 4 ผลการดำเนินงานและการทดสอบระบบ", "18"),
        ("    4.1 ผลการพัฒนาระบบส่วนติดต่อผู้ใช้งาน (UI Implementation)", "18"),
        ("    4.2 ผลการทดสอบเชิงฟังก์ชันและการตรวจสอบอัตโนมัติ (Automated Testing)", "24"),
        ("    4.3 ผลการประเมินประสิทธิภาพเชิงปฏิบัติการ (KPI Evaluation)", "25"),
        ("บทที่ 5 สรุปผล อภิปรายผล และข้อเสนอแนะ", "27"),
        ("    5.1 สรุปผลการดำเนินงาน", "27"),
        ("    5.2 การอภิปรายผลเชิงวิศวกรรม", "27"),
        ("    5.3 บทเรียนการพัฒนาและการแก้ปัญหาเชิงเทคนิค (Lessons Learned)", "28"),
        ("    5.4 ข้อเสนอแนะสำหรับการพัฒนาต่อยอด", "29"),
        ("เอกสารอ้างอิง (References)", "30"),
        ("ภาคผนวก", "32"),
        ("ประวัติผู้จัดทำ", "33")
    ]

    table_toc = doc.add_table(rows=len(toc_items), cols=2)
    for idx, (title, pg) in enumerate(toc_items):
        r_cells = table_toc.rows[idx].cells
        r_cells[0].text = ""
        r_cells[1].text = ""
        p0 = r_cells[0].paragraphs[0]
        p1 = r_cells[1].paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        is_bold = "บทที่" in title or "บทคัดย่อ" in title or "เอกสารอ้างอิง" in title or "สารบัญ" in title or "กิตติกรรมประกาศ" in title
        add_run(p0, title, size=15, bold=is_bold)
        add_run(p1, pg, size=15, bold=is_bold)

    doc.add_page_break()

    # 2.6 สารบัญตาราง & สารบัญภาพ (List of Tables & Figures)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "สารบัญตาราง", size=20, bold=True)

    lot_items = [
        ("ตารางที่ 2.1 ตารางเปรียบเทียบเทคโนโลยีที่เลือกใช้ในโครงงาน", "9"),
        ("ตารางที่ 3.1 ตารางแผนการดำเนินงานและไทม์ไลน์การพัฒนา 4 เดือน (Gantt Chart)", "11"),
        ("ตารางที่ 3.2 โครงสร้างตารางฐานข้อมูลหลักของระบบ QSMS บน Supabase (Data Dictionary)", "15"),
        ("ตารางที่ 4.1 สรุปผลการรันชุดทดสอบอัตโนมัติผ่าน Vitest (131 เคสการทดสอบ)", "24"),
        ("ตารางที่ 4.2 ตารางเปรียบเทียบประสิทธิภาพการทำงานก่อนและหลังนำระบบไปใช้จริง (KPIs)", "25")
    ]
    table_lot = doc.add_table(rows=len(lot_items), cols=2)
    for idx, (title, pg) in enumerate(lot_items):
        r_cells = table_lot.rows[idx].cells
        r_cells[0].text = ""
        r_cells[1].text = ""
        p0 = r_cells[0].paragraphs[0]
        p1 = r_cells[1].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p0, title, size=15, bold=False)
        add_run(p1, pg, size=15, bold=False)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, "สารบัญภาพ", size=20, bold=True)

    lof_items = [
        ("ภาพที่ 3.1 แผนผังสถาปัตยกรรมระบบโดยรวมและการไหลของข้อมูล (System Architecture & Data Flow)", "12"),
        ("ภาพที่ 3.2 แผนภาพความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram: ERD)", "14"),
        ("ภาพที่ 4.1 หน้าจอศูนย์รวมแอปพลิเคชันหลัก (Workspace Portal Shell)", "18"),
        ("ภาพที่ 4.2 หน้าต่างฟอร์มบันทึกเคส Rework พร้อม Two-Way Autofill และการตรวจสอบความถูกต้อง", "19"),
        ("ภาพที่ 4.3 แผงตรวจทานเอกสารข้อกำหนดผลิตภัณฑ์และสูตรมาตรฐานการผลิต Split View 55/45 พร้อมเครื่องมือหมุน PDF", "21"),
        ("ภาพที่ 4.4 หน้าต่างสนทนากับระบบปัญญาประดิษฐ์ DocAI RAG ('น้องผึ้งพา')", "22"),
        ("ภาพที่ 4.5 หน้าต่างจำลองไฟล์ส่งออก Excel Spreadsheet (.xlsx) พร้อมฝังรูปภาพในเซลล์", "23"),
        ("ภาพที่ 4.6 สไลด์คู่มือนำเสนอระบบ (Presentation Deck) และระบบ Direct-Manipulation Zoom/Pan", "23")
    ]
    table_lof = doc.add_table(rows=len(lof_items), cols=2)
    for idx, (title, pg) in enumerate(lof_items):
        r_cells = table_lof.rows[idx].cells
        r_cells[0].text = ""
        r_cells[1].text = ""
        p0 = r_cells[0].paragraphs[0]
        p1 = r_cells[1].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p0, title, size=15, bold=False)
        add_run(p1, pg, size=15, bold=False)

    # =============================================================
    # SECTION 3: CHAPTER 1 (Arabic numerals starting at 1)
    # =============================================================
    s_ch1 = doc.add_section()
    setup_section(s_ch1, is_first_page_diff=True)
    s_ch1.header.is_linked_to_previous = False
    add_page_number_to_header(s_ch1.header)
    set_section_pgnum(s_ch1, fmt='decimal', start=1)

    add_heading_chapter(doc, "บทที่ 1", "บทนำ (Introduction)")

    add_heading_1(doc, "1.1 ความเป็นมาและความสำคัญของปัญหา")
    add_body_p(doc, "ในอุตสาหกรรมการผลิตและการบรรจุภัณฑ์ การรักษามาตรฐานคุณภาพผลิตภัณฑ์ (Quality Assurance) และการบริหารจัดการสินค้าที่ไม่เป็นไปตามข้อกำหนด (Non-conforming Products) ถือเป็นหัวใจสำคัญในการลดความสูญเปล่า (Waste Reduction) และควบคุมต้นทุนการดำเนินงานขององค์กร เมื่อตรวจพบความผิดปกติในกระบวนการบรรจุ เช่น ผลิตภัณฑ์รั่วซึม (Leakage), ฝาบรรจุภัณฑ์ปิดไม่สนิท, ฉลากชำรุด หรือพบการปนเปื้อน สินค้าดังกล่าวจะถูกคัดแยกเข้าสู่กระบวนการทำใหม่ (Rework Workflow) เพื่อแก้ไขให้อยู่ในเกณฑ์มาตรฐาน")

    add_body_p(doc, "จากการศึกษาสภาพการปฏิบัติงานจริงในสายการผลิต พบว่ากระบวนการจัดการงาน Rework แบบดั้งเดิมที่ใช้เอกสารกระดาษและสเปรดชีตแบบกระจายตัว ประสบปัญหาสำคัญ 3 ประการ:")
    add_bullet_p(doc, "1. ความล่าช้าและข้อผิดพลาดในการบันทึกข้อมูล:", "การจดบันทึกด้วยมือทำให้รหัสสินค้า (Item Code / Item Number) และข้อมูลล็อตสินค้าคลาดเคลื่อน ส่งผลให้การสืบค้นประวัติย้อนกลับ (Traceability) ไม่สามารถทำได้แบบเรียลไทม์", level=1)
    add_bullet_p(doc, "2. การขาดหลักฐานเชิงประจักษ์ (Lack of Evidence Integrity):", "สินค้าที่เสียหายมักไม่มีการบันทึกภาพถ่ายหลักฐานอย่างเป็นระบบ หรือภาพถ่ายกระจัดกระจายอยู่ในแอปพลิเคชันส่งข้อความส่วนตัว ทำให้ไม่สามารถตรวจสอบยืนยันสภาพความเสียหายจริงในขั้นตอนการอนุมัติได้", level=1)
    add_bullet_p(doc, "3. ความยากลำบากในการเข้าถึงเอกสารข้อกำหนดผลิตภัณฑ์และสูตรมาตรฐานการผลิต:", "เอกสารแบบร่างข้อกำหนดผลิตภัณฑ์ของลูกค้า (Customer Product Drawings) และใบกำหนดสูตรและมาตรฐานการผลิตภายใน (Internal Manufacturing Master Sheets) มีปริมาณมากและจัดเก็บในรูปแบบเอกสารสแกน ส่งผลให้พนักงานต้องใช้เวลาเฉลี่ยถึง 18 นาทีต่อครั้งในการค้นหาข้อกำหนดทางเทคนิคและสูตรการผลิตที่ถูกต้อง", level=1)

    add_body_p(doc, "เพื่อแก้ไขปัญหาดังกล่าว ผู้จัดทำจึงได้พัฒนา 'ระบบบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์และคลังเอกสารข้อกำหนดมาตรฐานการผลิตอัจฉริยะ (QSMS: Quality & Rework Management System)' โดยผสานสถาปัตยกรรม Next.js App Router, ฐานข้อมูลเวกเตอร์ Supabase pgvector, เทคโนโลยีปัญญาประดิษฐ์ Multimodal OCR สำหรับวิเคราะห์เอกสารแบบร่างและสูตรการผลิต ตลอดจนระบบสืบค้นความรู้อัจฉริยะ (DocAI RAG) เข้าด้วยกันอย่างสมบูรณ์")

    add_heading_1(doc, "1.2 วัตถุประสงค์ของโครงงาน")
    add_bullet_p(doc, "1.", "เพื่อออกแบบและพัฒนาระบบเว็บแอปพลิเคชันบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์ (Rework Management System) แบบเรียลไทม์ พร้อมระบบตรวจสอบรหัสสินค้าสองทางอัตโนมัติ (Two-Way Autofill & Verification) และการรักษาความสมบูรณ์ของภาพถ่ายหลักฐาน (Evidence Integrity)")
    add_bullet_p(doc, "2.", "เพื่อพัฒนาระบบวิเคราะห์และสกัดข้อมูลจากเอกสารแบบร่างข้อกำหนดผลิตภัณฑ์ (Customer Drawings) และใบกำหนดสูตรมาตรฐานการผลิต (Manufacturing Master Sheets) ในรูปแบบ PDF ด้วยเทคโนโลยี Gemini Multimodal OCR แบบแยกสกีมาโครงสร้างข้อมูล (Decoupled Document Schemas)")
    add_bullet_p(doc, "3.", "เพื่อพัฒนาระบบถามตอบคู่มือเทคนิคและแนวทางการแก้ไขงานอัจฉริยะ (DocAI RAG) ด้วยการผสานฐานข้อมูลเวกเตอร์ Supabase pgvector ร่วมกับโมเดลการแปลงเวกเตอร์ความหมาย Jina AI Embeddings และระบบค้นหาแบบผสมผสาน (Hybrid Search)")
    add_bullet_p(doc, "4.", "เพื่อพัฒนาระบบส่งออกรายงานสเปรดชีต Excel (.xlsx) ที่ฝังรูปภาพหลักฐานความละเอียดสูงลงในเซลล์ตารางโดยตรง พร้อมระบบสรุปสถิติแดชบอร์ดเพื่อการตัดสินใจเชิงคุณภาพ")
    add_bullet_p(doc, "5.", "เพื่อพัฒนาระบบคู่มือนำเสนอและสภาพแวดล้อมจำลองเชิงโต้ตอบ (Interactive Presentation Deck & Sandbox) 20 สไลด์ พร้อมกลไกควบคุมมุมมองการนำเสนอแบบ Direct-Manipulation 1:1 Zoom and Pan Engine")

    add_heading_1(doc, "1.3 ขอบเขตของโครงงาน (Scope)")
    add_heading_2(doc, "1.3.1 ขอบเขตด้านการทำงาน (5 โมดูลหลัก)")
    add_bullet_p(doc, "Workspace Portal & Auth Shell:", "ระบบยืนยันตัวตนระดับเซิร์ฟเวอร์ (HTTP-Only Cookie Session), หน้าจอศูนย์รวมแอปพลิเคชัน (Portal Shell) และการแบ่งสิทธิ์ RBAC (Admin/QSMS และ Operator/WFG)", level=1)
    add_bullet_p(doc, "Rework Module:", "ระบบบริหารจัดการงานแก้ไขสินค้าแบบครบวงจร ประกอบด้วยฟอร์มสร้างเคส Multi-Item, Two-Way Autofill, การเชื่อมโยงสินค้าเปื้อน-รั่ว (Cross-Item Link), กล่องแนบไฟล์ OR, ตารางภาพรวมพร้อมการจัดลำดับคิวงาน, แผงจัดการเคสพร้อม Item Accordion และ Dynamic Auto-Status, แดชบอร์ดสรุปสถิติและวิเคราะห์ Defect Pareto, และระบบบันทึกด่วนหน้าไลน์บนมือถือ (Mobile Fast-Track)", level=1)
    add_bullet_p(doc, "Drawing & Master Storage Module:", "คลังเอกสารข้อกำหนดผลิตภัณฑ์ของลูกค้าและสูตรมาตรฐานการผลิต, ระบบ AI OCR สกัด 7-8 ฟิลด์, แผงตรวจทาน Split View 55/45, และระบบหมุน PDF 360 องศา", level=1)
    add_bullet_p(doc, "DocAI RAG Module:", "แชทบอท 'น้องผึ้งพา' สืบค้นคู่มือเทคนิคด้วย Hybrid Search, Markdown Streaming, และ Function Calling สถิติสด", level=1)
    add_bullet_p(doc, "Guide & Presentation Module:", "คู่มือนำเสนอ 20 สไลด์ สไตล์ Apple Liquid Glass พร้อม Live Sandboxes และระบบควบคุม Zoom/Pan 1:1", level=1)

    add_heading_2(doc, "1.3.2 ขอบเขตด้านเทคโนโลยี")
    add_bullet_p(doc, "เทคโนโลยีฝั่งหน้าบ้าน (Frontend):", "Next.js 16.2 App Router, React 19, Tailwind CSS v4, Motion (Framer Motion), Radix UI", level=1)
    add_bullet_p(doc, "เทคโนโลยีฝั่งหลังบ้าน (Backend):", "Next.js API Route Handlers, TypeScript 5.8 (Strict Type Safety, No Any)", level=1)
    add_bullet_p(doc, "ฐานข้อมูลและการจัดเก็บ (Database & Storage):", "Supabase PostgreSQL + pgvector, Cloudinary CDN (Unsigned Direct Upload)", level=1)
    add_bullet_p(doc, "ปัญญาประดิษฐ์ (AI & LLM):", "Google Gemini 3.1 Flash/Lite (Vision OCR), Jina AI Embeddings v5 (768 Dimensions)", level=1)
    add_bullet_p(doc, "การทดสอบ (Testing):", "Vitest (Unit/Integration Tests 131 เคส) และ Playwright (End-to-End Tests)", level=1)

    add_heading_1(doc, "1.4 นิยามศัพท์เฉพาะ (Operational Definitions)")
    add_bullet_p(doc, "Case ID:", "รหัสประจำใบงาน Rework ที่ระบบออกให้อัตโนมัติและไม่สามารถแก้ไขได้ เช่น RW012-2026 หรือ RT012-2026", level=1)
    add_bullet_p(doc, "Two-Way Autofill:", "ระบบตรวจสอบสินค้าอัตโนมัติ เมื่อกรอก Item Number หรือ Item Code ระบบจะค้นหาและดึงข้อมูลจาก Item Master มากรอกให้อัตโนมัติ", level=1)
    add_bullet_p(doc, "Evidence Integrity:", "กฎข้อบังคับที่กำหนดให้ทุกรายการสินค้าชำรุดต้องมีภาพถ่ายหลักฐานอย่างน้อย 1 ภาพ โดยบันทึกแบบ Atomic Transaction", level=1)
    add_bullet_p(doc, "Boxes Per Pallet Normalization:", "กฎการแปลงค่าจำนวนกล่องต่อพาเลท หากพบคำว่า 'ตามความเหมาะสม' ระบบจะคงค่าข้อความไว้โดยไม่สุ่มเดาตัวเลข", level=1)
    add_bullet_p(doc, "Direct-Manipulation Zoom/Pan:", "ระบบควบคุมมุมมองสไลด์ในพิกัด 1:1 โดยตัดการหน่วงของ CSS Transition ระหว่างลากเมาส์ เพื่อให้ตอบสนองทันที", level=1)

    add_heading_1(doc, "1.5 ประโยชน์ที่คาดว่าจะได้รับ")
    add_bullet_p(doc, "1.", "ลดระยะเวลาในการกรอกและตรวจสอบข้อมูลเคส Rework ลงได้มากกว่า 60%")
    add_bullet_p(doc, "2.", "ขจัดข้อผิดพลาดของรหัสสินค้าและล็อตสินค้าในระบบให้เหลือ 0%")
    add_bullet_p(doc, "3.", "ลดเวลาในการสืบค้นเอกสารข้อกำหนดและคู่มือเทคนิคโรงงานลงได้มากกว่า 90%")
    add_bullet_p(doc, "4.", "ยกระดับความโปร่งใสและตรวจสอบย้อนกลับได้ของกระบวนการควบคุมคุณภาพในองค์กร")

    # =============================================================
    # SECTION 4: CHAPTER 2 (Arabic continue)
    # =============================================================
    s_ch2 = doc.add_section()
    setup_section(s_ch2, is_first_page_diff=True)
    s_ch2.header.is_linked_to_previous = False
    add_page_number_to_header(s_ch2.header)

    add_heading_chapter(doc, "บทที่ 2", "ทฤษฎีและงานวิจัยที่เกี่ยวข้อง (Literature Review)")

    add_heading_1(doc, "2.1 สถาปัตยกรรมเว็บแอปพลิเคชันยุคใหม่ (Modern Web Architecture)")
    add_body_p(doc, "ในการพัฒนาระบบสารสนเทศระดับองค์กรยุคปัจจุบัน สถาปัตยกรรม Next.js App Router (เวอร์ชัน 16) ได้รับการยอมรับอย่างกว้างขวางเนื่องจากสามารถรวมขีดความสามารถของ Server-Side Rendering (SSR), Server Components และ Client Components เข้าด้วยกันได้อย่างมีประสิทธิภาพ")

    add_body_p(doc, "ระบบ QSMS ประยุกต์ใช้สถาปัตยกรรม Hybrid Next.js + React 19 ควบคู่กับระเบียบวิธี Feature-Sliced Design (FSD) ซึ่งเป็นการแบ่งสัดส่วนโค้ดออกเป็นโมดูลอิสระ (Module Slices) ได้แก่ auth, rework, drawings, rag, guide และ platform ข้อดีของสถาปัตยกรรมนี้คือการลดความผูกมัดของโค้ด (Loose Coupling) ป้องกันปัญหาการพึ่งพากันแบบวงกลม (Circular Dependencies) และเอื้อต่อการเขียนชุดทดสอบอัตโนมัติ")

    add_heading_1(doc, "2.2 ระบบความปลอดภัยและการจัดการสิทธิ์การใช้งาน (Security & RBAC)")
    add_body_p(doc, "ระบบรักษาความปลอดภัยของ QSMS ยึดหลัก Defense-in-Depth และมาตรฐาน OWASP โดยมีองค์ประกอบสำคัญ 2 ส่วน:")
    add_bullet_p(doc, "Server-State Authentication via HTTP-Only Cookies:", "ระบบยกเลิกการเก็บ JWT Token ใน localStorage บนเบราว์เซอร์ เพื่อป้องกันการถูกโจมตีด้วยช่องโหว่ Cross-Site Scripting (XSS) โดยเปลี่ยนมาใช้คุกกี้ที่ปลอดภัย (HTTP-Only, Secure, SameSite=Lax) จัดการผ่าน Next.js Server Boundary (`src/lib/serverAuth.ts`)", level=1)
    add_bullet_p(doc, "Role-Based Access Control (RBAC):", "ระบบแบ่งสิทธิ์ผู้ใช้งานออกเป็น 2 ระดับหลัก ได้แก่ (1) Admin/QSMS ซึ่งมีสิทธิ์เต็มในการเข้าถึงแดชบอร์ดและโหมดแก้ไขข้อมูลเคส และ (2) Operator/WFG ซึ่งจำกัดสิทธิ์เฉพาะการบันทึกเคสใหม่และการอัปเดตยอดกล่องที่ทำเสร็จ โดยซ่อนฟังก์ชันด้านต้นทุนและแดชบอร์ด", level=1)

    add_heading_1(doc, "2.3 เทคโนโลยีปัญญาประดิษฐ์ประมวลผลเอกสารหลายรูปแบบ (Multimodal AI & OCR)")
    add_body_p(doc, "การสกัดข้อมูลจากแบบร่างข้อกำหนดผลิตภัณฑ์ของลูกค้า (Customer Drawings) และใบกำหนดสูตรมาตรฐานการผลิต (Master Sheets) ซึ่งเป็นไฟล์ PDF สแกน มีความท้าทายเนื่องจากโครงสร้างตารางและตัวอักษรมีความซับซ้อน ระบบ QSMS เลือกใช้โมเดล Google Gemini Multimodal Vision API (`gemini-3.1-flash`) ร่วมกับเทคนิค Structured Outputs (JSON Schema Enforcement) ซึ่งบังคับให้โมเดลส่งผลลัพธ์กลับมาเป็นโครงสร้างข้อมูล JSON ตามสเปกที่กำหนดโดยตรง")

    add_body_p(doc, "นอกจากนี้ เพื่อป้องกันปัญหาความไม่เสถียรในช่วงเวลาที่มีผู้ใช้งานหนาแน่น (503 Service Unavailable Spikes) ระบบได้พัฒนาสถาปัตยกรรม Fallback Pipeline อัตโนมัติ โดยหากโมเดลหลัก `gemini-3.1-flash` ไม่สามารถตอบสนองได้ ระบบจะสลับไปยัง `gemini-3.1-flash-lite` และ `gemini-2.0-flash` ตามลำดับ")

    add_heading_1(doc, "2.4 สถาปัตยกรรม Retrieval-Augmented Generation (RAG)")
    add_body_p(doc, "สถาปัตยกรรม RAG ทำหน้าที่เชื่อมโยงคลังเอกสารคู่มือโรงงานเข้ากับโมเดลภาษาขนาดใหญ่ เพื่อให้ระบบสามารถตอบคำถามทางเทคนิคได้อย่างถูกต้อง แม่นยำ และอ้างอิงเอกสารจริงได้ โดยมีกระบวนการสำคัญดังนี้:")
    add_bullet_p(doc, "Document Parsing & Ingestion:", "แปลงหน้าเอกสาร PDF เป็นภาพความละเอียดสูงด้วย `pdfjs-dist` จากนั้นใช้ Gemini Vision ถอดเนื้อหาเป็น Markdown และแบ่งท่อนข้อมูล (Chunking) ตามโครงสร้างหัวข้อ", level=1)
    add_bullet_p(doc, "Vector Embeddings:", "แปลงท่อนข้อความ Markdown เป็นเวกเตอร์ความหมายขนาด 768 มิติด้วยโมเดล `jina-embeddings-v5-text-small` และบันทึกลงในตาราง `rag_document_chunks` บน Supabase pgvector", level=1)
    add_bullet_p(doc, "Hybrid Search Engine:", "เมื่อผู้ใช้ถามคำถาม ระบบจะทำการค้นหาแบบผสมผสานระหว่าง Cosine Similarity Search บนเวกเตอร์ และ Full-Text Search บนคีย์เวิร์ดภาษาไทย ผ่าน RPC Function ของ PostgreSQL", level=1)
    add_bullet_p(doc, "Function Calling Integration:", "เมื่อผู้ใช้สอบถามสถิติเคส Rework ระบบเปิดให้โมเดลเรียกฟังก์ชัน `get_rework_statistics` เพื่อดึงข้อมูลจริงจากตาราง `rework_cases` มาสังเคราะห์คำตอบแบบเรียลไทม์", level=1)

    add_heading_1(doc, "2.5 ตารางสรุปเปรียบเทียบเทคโนโลยีที่เลือกใช้ในโครงงาน")
    
    add_table_caption(doc, "ตารางที่ 2.1 ตารางเปรียบเทียบเทคโนโลยีที่เลือกใช้ในโครงงาน")

    table_tech = doc.add_table(rows=6, cols=4)
    table_tech_widths = [Inches(1.5), Inches(1.8), Inches(1.5), Inches(1.8)]
    table_tech_headers = ["หมวดหมู่", "เทคโนโลยีที่เลือกใช้", "เทคโนโลยีทางเลือก", "เหตุผลทางวิศวกรรมที่เลือก"]
    table_tech_data = [
        ["Web Framework", "Next.js 16 + React 19", "Vite + Express SPA", "มี Server Boundary ในตัว จัดการ Auth Cookie และ API Proxy ได้ปลอดภัย"],
        ["Database & Vector", "Supabase (PostgreSQL + pgvector)", "MongoDB + Pinecone", "รวม Relational Data และ Vector Embeddings ไว้ในระบบเดียวกัน ลดความซับซ้อน"],
        ["AI Multimodal OCR", "Gemini 3.1 Flash (Structured JSON)", "Tesseract OCR / AWS Textract", "รองรับภาษาไทย เข้าใจโครงสร้างตารางวิศวกรรม และส่ง JSON schema ได้ตรง"],
        ["Vector Embeddings", "Jina Embeddings v5 (768 Dim)", "OpenAI text-embedding-3", "รองรับ Multilingual (ไทย-อังกฤษ) คุณภาพสูง ขนาดเวกเตอร์กะทัดรัด ประหยัด Storage"],
        ["Testing Framework", "Vitest + Playwright", "Jest + Cypress", "ความเร็วในการรันสูง รองรับ ESM สมบูรณ์แบบ และทดสอบ E2E ได้เสถียรบน Chromium"]
    ]
    style_academic_table(table_tech, table_tech_widths, table_tech_headers, table_tech_data)

    # =============================================================
    # SECTION 5: CHAPTER 3 (Arabic continue)
    # =============================================================
    s_ch3 = doc.add_section()
    setup_section(s_ch3, is_first_page_diff=True)
    s_ch3.header.is_linked_to_previous = False
    add_page_number_to_header(s_ch3.header)

    add_heading_chapter(doc, "บทที่ 3", "วิธีดำเนินการและการออกแบบระบบ (System Design)")

    add_heading_1(doc, "3.1 ระเบียบวิธีและแผนการดำเนินงาน 4 เดือน (Project Timeline & Methodology)")
    add_body_p(doc, "การพัฒนาระบบ QSMS ดำเนินการตามแนวคิดการพัฒนาซอฟต์แวร์แบบ Agile Software Development Lifecycle โดยแบ่งระยะเวลาการดำเนินงานออกเป็น 4 เดือน (พฤศจิกายน 2567 ถึง กุมภาพันธ์ 2568) ครอบคลุม 4 เฟสหลัก ดังแสดงในตาราง Gantt Chart ที่ 3.1:")

    add_table_caption(doc, "ตารางที่ 3.1 ตารางแผนการดำเนินงานและไทม์ไลน์การพัฒนา 4 เดือน (Gantt Chart)")

    table_gantt = doc.add_table(rows=9, cols=6)
    table_gantt_widths = [Inches(0.6), Inches(2.6), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9)]
    table_gantt_headers = ["ลำดับ", "กิจกรรมการดำเนินงาน (Activities)", "พ.ย. 67", "ธ.ค. 67", "ม.ค. 68", "ก.พ. 68"]
    table_gantt_data = [
        ["1", "ศึกษาปัญหา วิเคราะห์ Workflow Rework และรวบรวม Requirement", "■■■■", "", "", ""],
        ["2", "ออกแบบสถาปัตยกรรม Next.js 16 App Router และ Database Schemas", "■■", "■■", "", ""],
        ["3", "พัฒนา Core Rework Module, Two-Way Autofill & Cloudinary Upload", "■■", "■■■■", "", ""],
        ["4", "พัฒนาระบบ AI Multimodal OCR วิเคราะห์แบบร่างและใบกำหนดสูตรผลิต", "", "■■■■", "■■", ""],
        ["5", "พัฒนาระบบ DocAI RAG (pgvector + Jina Embeddings) & Function Calling", "", "", "■■■■", ""],
        ["6", "พัฒนาระบบส่งออกรายงาน Excel ฝังรูปภาพในเซลล์ & Presentation Deck", "", "", "", "■■■■"],
        ["7", "ทดสอบระบบอัตโนมัติ (Vitest 131 Tests) และประเมินผลระบบหน้างาน", "", "", "■■", "■■■■"],
        ["8", "จัดทำเอกสารคู่มือการใช้งานและจัดพิมพ์เล่มปริญญานิพนธ์ฉบับสมบูรณ์", "", "", "", "■■■■"]
    ]
    aligns_gantt = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    style_academic_table(table_gantt, table_gantt_widths, table_gantt_headers, table_gantt_data, align_cols=aligns_gantt)

    add_heading_1(doc, "3.2 สถาปัตยกรรมระบบโดยรวม (System Architecture & Data Flow)")
    add_body_p(doc, "สถาปัตยกรรมของ QSMS ถูกออกแบบตามหลักการแยกหน้าที่ (Separation of Concerns) แบ่งออกเป็น 3 เลเยอร์หลักตามมาตรฐานระบบสารสนเทศระดับองค์กร ดังแสดงในภาพที่ 3.1:")
    add_bullet_p(doc, "1. Client Presentation Layer (React 19 SPA):", "โครงสร้างโฟลเดอร์แบบ Feature-Sliced Design (FSD) ประกอบด้วย 5 โมดูลอิสระ (rework, drawings, rag, guide, platform) ทำหน้าที่บริหารจัดการ State ภายในหน้าจอและบีบอัดภาพก่อนส่งขึ้น CDN", level=1)
    add_bullet_p(doc, "2. Next.js 16 App Router Boundary & Server Services:", "ทำหน้าที่เป็นตัวกลางในการตรวจสอบสิทธิ์ (JWT Verification via HTTP-Only Secure Cookies), ซ่อน API Keys ความลับ, จัดการ Atomic Transactions, และทำหน้าที่เป็น Proxy ไปยัง Supabase, Cloudinary และ Google Gemini API", level=1)
    add_bullet_p(doc, "3. Persistence, Vector Store & Multimodal Intelligence Layer:", "ประกอบด้วย Supabase PostgreSQL สำหรับเก็บข้อมูลตารางความสัมพันธ์, Supabase pgvector สำหรับเก็บเวกเตอร์คู่มือ 768 มิติ, Cloudinary สำหรับเก็บไฟล์รูปภาพหลักฐาน และ Gemini LLM สำหรับการประมวลผล OCR และการตอบคำถาม", level=1)

    add_embedded_image(doc, arch_diag_path, width_in=6.0)
    add_figure_caption(doc, "ภาพที่ 3.1 แผนผังสถาปัตยกรรมระบบโดยรวมและการไหลของข้อมูล (System Architecture & Data Flow)", "ผู้จัดทำ, 2568")

    add_heading_1(doc, "3.3 การออกแบบโครงสร้างฐานข้อมูล (Entity-Relationship Diagram & Data Dictionary)")
    add_body_p(doc, "ระบบ QSMS จัดเก็บข้อมูลแบบผสมผสานระหว่างข้อมูลเชิงสัมพันธ์ (Relational Data) และข้อมูลเวกเตอร์ความหมาย (Vector Embeddings) โดยมีแผนภาพความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram: ERD) ดังแสดงในภาพที่ 3.2:")

    add_embedded_image(doc, erd_diag_path, width_in=6.0)
    add_figure_caption(doc, "ภาพที่ 3.2 แผนภาพความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram: ERD)", "ผู้จัดทำ, 2568")

    add_body_p(doc, "รายละเอียดของตารางฐานข้อมูลและพจนานุกรมข้อมูล (Data Dictionary) บน Supabase PostgreSQL สรุปได้ดังตารางที่ 3.2:")

    add_table_caption(doc, "ตารางที่ 3.2 โครงสร้างตารางฐานข้อมูลหลักของระบบ QSMS บน Supabase (Data Dictionary)")

    table_db = doc.add_table(rows=7, cols=4)
    table_db_widths = [Inches(1.8), Inches(1.5), Inches(1.8), Inches(1.8)]
    table_db_headers = ["ชื่อตาราง (Table Name)", "คีย์หลัก (Primary Key)", "คีย์นอก (Foreign Key)", "หน้าที่และการจัดเก็บข้อมูล"]
    table_db_data = [
        ["rework_cases", "id (UUID)", "created_by -> users.id", "จัดเก็บข้อมูลหัวใบงาน Rework (Case ID, แหล่งที่มา SFC/Customer, สถานะ, กำหนดเสร็จ)"],
        ["rework_items", "id (UUID)", "case_id -> rework_cases.id", "จัดเก็บรายการสินค้าชำรุด (Item Code, Item Number, จำนวน, ล็อต, เลขกล่อง, รูปภาพ, Defect Type)"],
        ["rework_master_items", "id (UUID)", "-", "จัดเก็บฐานข้อมูลมาสเตอร์สินค้าส่วนกลาง (Part Name, Oil Group, Pallet Type, Boxes/Pallet)"],
        ["drawings", "id (UUID)", "-", "จัดเก็บเอกสารแบบร่างข้อกำหนดและใบมาสเตอร์สูตรผลิต พร้อม Metadata ที่สกัดได้จาก Gemini Vision OCR"],
        ["rag_documents", "id (UUID)", "-", "จัดเก็บหัวข้อเอกสารคู่มือโรงงาน ประเภทเอกสาร และสถานะการประมวลผลเวกเตอร์"],
        ["rag_document_chunks", "id (UUID)", "document_id -> rag_documents.id", "จัดเก็บท่อนเนื้อหาคู่มือ Markdown พร้อมเวกเตอร์ Embeddings 768 มิติ (vector(768))"]
    ]
    style_academic_table(table_db, table_db_widths, table_db_headers, table_db_data)

    add_heading_1(doc, "3.4 การออกแบบขั้นตอนการทำงานและสมการทางวิศวกรรม (Core Equations & Algorithms)")
    
    add_heading_2(doc, "3.4.1 สมการคำนวณสถานะเคสอัตโนมัติ (Dynamic Auto-Status Lifecycle)")
    add_body_p(doc, "ระบบใช้อัลกอริทึมคำนวณสถานะอัตโนมัติจากผลรวมจำนวนกล่องที่ผลิตเสร็จจริง (completedBoxes) เทียบกับยอดรวมทั้งหมด (totalBoxes) ตามสมการที่ (3.1):")
    
    add_equation_block(doc, "Status = { Pending  if  completedBoxes = 0 ; In-Progress  if  0 < completedBoxes < totalBoxes ; Completed  if  completedBoxes >= totalBoxes }", "3.1")

    add_heading_2(doc, "3.4.2 สมการควบคุมมุมมองนำเสนอแบบ Direct-Manipulation (1:1 Zoom & Pan Engine)")
    add_body_p(doc, "ในการแสดงผลสไลด์คู่มือนำเสนอบน Canvas ความละเอียด 2560x1440 พิกเซล การคำนวณตำแหน่งการเลื่อนหน้าจอจะใช้ Screen-Space Translation ร่วมกับ Scale Matrix ตามสมการที่ (3.2):")
    
    add_equation_block(doc, "Transform = translate3d(Δx, Δy, 0px) · scale(S_fit · Z_user)", "3.2")
    
    add_body_p(doc, "โดยในขณะที่ผู้ใช้งานกำลังคลิกลากหน้าจอ (Dragging State) ระบบจะปลดคลาส CSS Transition ออก (transition-none) เพื่อให้การเคลื่อนไหวตอบสนองต่อพิกัดเมาส์ทันทีโดยไม่มีอาการกระตุกหรือหน่วงเวลา")

    # =============================================================
    # SECTION 6: CHAPTER 4 (Arabic continue)
    # =============================================================
    s_ch4 = doc.add_section()
    setup_section(s_ch4, is_first_page_diff=True)
    s_ch4.header.is_linked_to_previous = False
    add_page_number_to_header(s_ch4.header)

    add_heading_chapter(doc, "บทที่ 4", "ผลการดำเนินงานและการทดสอบระบบ (Results & Testing)")

    add_heading_1(doc, "4.1 ผลการพัฒนาระบบส่วนติดต่อผู้ใช้งาน (UI Implementation)")
    add_body_p(doc, "ระบบ QSMS ได้รับการพัฒนาจนเสร็จสมบูรณ์และสามารถทำงานได้ครบทั้ง 5 โมดูลหลัก โดยมีผลการทำงานและภาพหน้าจอจริงของระบบดังนี้:")
    
    # 4.1 Workspace Portal Shell
    add_bullet_p(doc, "1. Workspace Portal Shell:", "หน้าจอหลักรวมศูนย์แอปพลิเคชันที่มาพร้อม Live Preview Analytics สำหรับผู้ใช้ทั่วไป และระบบสลับโมดูลได้อย่างราบรื่น ดังแสดงในภาพที่ 4.1", level=1)
    add_embedded_image(doc, ui_workspace_path, width_in=5.8)
    add_figure_caption(doc, "ภาพที่ 4.1 หน้าจอศูนย์รวมแอปพลิเคชันหลัก (Workspace Portal Shell)", "ผู้จัดทำ, 2568")

    # 4.2 Rework Module
    add_bullet_p(doc, "2. Rework Module (ระบบบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์):", "โมดูลหลักสำหรับควบคุมและติดตามวงจรชีวิตของงาน Rework ทั้งหมดในสายการผลิตและงานส่งคืนจากลูกค้า ประกอบด้วย 5 ฟังก์ชันย่อยสำคัญ ดังแสดงในภาพที่ 4.2:", level=1)
    add_embedded_image(doc, ui_rework_path, width_in=5.8)
    add_figure_caption(doc, "ภาพที่ 4.2 หน้าต่างฟอร์มบันทึกเคส Rework พร้อม Two-Way Autofill และการตรวจสอบความถูกต้อง", "ผู้จัดทำ, 2568")

    add_body_p(doc, "จากภาพที่ 4.2 โมดูล Rework ได้รับการออกแบบให้รองรับกระบวนการแก้ไขผลิตภัณฑ์อย่างครบวงจร โดยครอบคลุม 5 ฟังก์ชันย่อยเชิงปฏิบัติการ ได้แก่:")
    
    add_bullet_p(doc, "ก. ฟอร์มบันทึกเคสใหม่และการตรวจสอบสินค้าสองทาง (Add Case & Two-Way Autofill):", "ระบบรองรับการออกรหัสเคสอัตโนมัติแบบไฮบริด (เช่น RW012-2026 สำหรับ SFC และ RT012-2026 สำหรับ Customer) ซึ่งมีคุณสมบัติไม่สามารถซ้ำซ้อนหรือแก้ไขได้ (Immutable) เมื่อผู้ใช้กรอกรหัส Item Number (สูตรการผลิต) หรือ Item Code (บาร์โค้ด) ระบบจะค้นหาและดึงข้อมูลจากฐานข้อมูล Item Master มาเติมให้อัตโนมัติ (Debounce 600ms) พร้อมแสดงสถานะการตรวจสอบผ่าน Verification Lifecycle (Idle ➔ Checking ➔ Verified / New / Conflict) นอกจากนี้ยังมีระบบเชื่อมโยงสินค้าเปื้อนกับสินค้าที่รั่วในเคสเดียวกัน (Cross-Item Link บันทึกลง linkedSourceId), ช่องแนบเอกสารอ้างอิงของลูกค้า OR สูงสุด 2 ไฟล์ (.xlsx, .pdf, .png) และการรักษาความสมบูรณ์ของหลักฐาน (Evidence Integrity) บังคับแนบภาพถ่ายชำรุดอย่างน้อย 1 ภาพ บีบอัดภาพฝั่งไคลเอนต์ (300KB) และบันทึกแบบธุรกรรมสมบูรณ์ (Atomic Transaction)", level=2)

    add_bullet_p(doc, "ข. หน้าจอภาพรวมและการจัดลำดับคิวงาน (Overall View & Priority Queue):", "ตารางแสดงรายการเคสทั้งหมดแบบเรียลไทม์ พร้อมระบบจัดเรียงลำดับความสำคัญ โดยดันเคสสถานะ Pending (รอดำเนินการ) ขึ้นบนสุดเสมอ เพื่อป้องกันไม่ให้งานตกค้างถูกละเลย จากนั้นจึงตามด้วย In-Progress และ Completed พร้อมระบบตัวกรองขั้นสูง (Advanced Filters) ตามแหล่งที่มา สาเหตุ ผู้รับผิดชอบ และป้ายเตือน Badge สีแดง 'ขาดไฟล์ OR' ในกรณีที่เคสของลูกค้า OR ยังไม่ได้แนบไฟล์เอกสาร", level=2)

    add_bullet_p(doc, "ค. แผงจัดการและอัปเดตงาน Rework (Case Update View & Item Accordion):", "หน้าต่างจัดการเคสรองรับการคำนวณสถานะอัตโนมัติ (Dynamic Auto-Status Lifecycle) จากจำนวนกล่องที่ผลิตเสร็จจริงเทียบกับยอดรวมทั้งหมด (Pending ➔ In-Progress ➔ Completed) โดยไม่ต้องเลือกสถานะเอง มีช่องบันทึกอุปสรรคหน้างาน (ขาดกล่อง, ขาดแกลลอน, ขาดน้ำมัน) เพื่อรายงานหัวหน้างานทันที พร้อมระบบพับเก็บรายการสินค้าเป็นค่าเริ่มต้น (Default Folded Item Accordion) และปุ่มลัดสลับ 'ขยาย/พับทั้งหมด' ร่วมกับแถบ Header 2 ชั้นแบบ Responsive ป้องกันปุ่มคำสั่งซ้อนทับกันบนอุปกรณ์พกพา รวมถึงโหมดแก้ไขข้อมูลระดับผู้ดูแล (Admin Edit Mode) พร้อมระบบแจ้งเตือนยืนยันป้องกันข้อมูลสูญหาย (Unsaved Changes Warning Guard)", level=2)

    add_bullet_p(doc, "ง. แดชบอร์ดสรุปสถิติเชิงปฏิบัติการ (Rework Operations & Defect Dashboard):", "ระบบประมวลผลและแสดงผลตัวชี้วัดประสิทธิภาพหลัก (KPI Summary Cards) ได้แก่ ยอดเคสรวม, จำนวนกล่องทั้งหมด, อัตราความสำเร็จ (Completion Rate %), และระยะเวลาเฉลี่ยในการปิดเคส (Lead Time) ควบคู่กับแผนภูมิวิเคราะห์สาเหตุชำรุดยอดนิยม (Defect Pareto Analysis) จำแนกตามไลน์การผลิตและกลุ่มผลิตภัณฑ์ ช่วยให้วิศวกรควบคุมคุณภาพสามารถระบุปัญหาคอขวดในกระบวนการผลิตได้อย่างแม่นยำ", level=2)

    add_bullet_p(doc, "จ. ระบบบันทึกด่วนหน้าไลน์ผลิตบนอุปกรณ์พกพา (Mobile Fast-Track UI & Local Draft Queue):", "ส่วนต่อประสานแบบ Mobile-First Card Layout ที่ได้รับการปรับแต่งให้เหมาะกับหน้าจอสมาร์ตโฟน (เช่น iPhone 13 ความกว้าง 390px) ขยายชื่อสินค้าเต็มความกว้าง จัดกลุ่มยอดผลิตและหลอดความคืบหน้าเข้ามุมล่างอย่างเป็นระเบียบ พร้อมระบบคิวข้อมูลสำรองในเครื่อง (Local Storage Draft Queue) สำหรับบันทึกข้อมูลร่างเคสต่อเนื่องแม้ในจุดที่สัญญาณอินเทอร์เน็ตหน้าไลน์ผลิตไม่เสถียร และปุ่มผู้ช่วยอัจฉริยะลอยตัว DocAI FAB Icon ที่มุมขวาล่างที่ไม่บดบังปุ่มเปลี่ยนหน้า", level=2)

    # 4.3 Drawing & Master Storage
    add_bullet_p(doc, "3. Drawing & Master Storage (คลังข้อกำหนดผลิตภัณฑ์และสูตรมาตรฐานการผลิต):", "คลังจัดเก็บเอกสารแบบร่างผลิตภัณฑ์ของลูกค้าและใบกำหนดสูตรมาตรฐานการผลิตพร้อม AI OCR อัตโนมัติ, การจัดแสดงแบบฟอร์มแยกประเภท Decoupled Forms, และแผงตรวจทานเอกสารแบบ Split View 55/45 พร้อมเครื่องมือหมุน PDF ดังแสดงในภาพที่ 4.3", level=1)
    add_embedded_image(doc, ui_drawings_path, width_in=5.8)
    add_figure_caption(doc, "ภาพที่ 4.3 แผงตรวจทานเอกสารข้อกำหนดผลิตภัณฑ์และสูตรมาตรฐานการผลิต Split View 55/45 พร้อมเครื่องมือหมุน PDF", "ผู้จัดทำ, 2568")

    # 4.4 DocAI RAG Chatbot
    add_bullet_p(doc, "4. DocAI RAG Chatbot ('น้องผึ้งพา'):", "แผงแชทถามตอบข้อสงสัยคู่มือเทคนิคภาษาไทย รองรับการแสดงผล Markdown, รูปภาพประกอบคู่มือ, Suggestion Chips, และการเรียกดูสถิติสดย้อนหลัง ดังแสดงในภาพที่ 4.4", level=1)
    add_embedded_image(doc, ui_rag_path, width_in=5.8)
    add_figure_caption(doc, "ภาพที่ 4.4 หน้าต่างสนทนากับระบบปัญญาประดิษฐ์ DocAI RAG ('น้องผึ้งพา')", "ผู้จัดทำ, 2568")

    # 4.5 Excel Export Preview
    add_bullet_p(doc, "5. Excel Export & Report Modal:", "หน้าต่างจำลองและส่งออกรายงาน Excel (.xlsx) ที่ฝังรูปภาพหลักฐานความละเอียดสูงลงในเซลล์โดยตรง (Row Height 120px) สไตล์ Microsoft Excel Ribbon ดังแสดงในภาพที่ 4.5", level=1)
    add_embedded_image(doc, ui_excel_path, width_in=5.8)
    add_figure_caption(doc, "ภาพที่ 4.5 หน้าต่างจำลองไฟล์ส่งออก Excel Spreadsheet (.xlsx) พร้อมฝังรูปภาพในเซลล์", "ผู้จัดทำ, 2568")

    # 4.6 Presentation Deck
    add_bullet_p(doc, "6. Interactive Presentation Deck:", "สไลด์นำเสนอ 20 สไลด์ สไตล์ Apple Liquid Glass พร้อม Live Interactive Sandbox สำหรับการทดลองใช้งานเสมือนจริง และระบบควบคุมมุมมอง Direct-Manipulation Zoom/Pan 1:1 ดังแสดงในภาพที่ 4.6", level=1)
    add_embedded_image(doc, ui_guide_path, width_in=5.8)
    add_figure_caption(doc, "ภาพที่ 4.6 สไลด์คู่มือนำเสนอระบบ (Presentation Deck) และระบบ Direct-Manipulation Zoom/Pan", "ผู้จัดทำ, 2568")

    add_heading_1(doc, "4.2 ผลการทดสอบเชิงฟังก์ชันและการตรวจสอบอัตโนมัติ (Automated Testing)")
    add_body_p(doc, "เพื่อรับประกันความเสถียรและความถูกต้องตามหลักวิศวกรรมซอฟต์แวร์ ระบบได้ผ่านการทดสอบอัตโนมัติอย่างเข้มงวดผ่านชุดทดสอบ Vitest และ TypeScript Compiler ดังแสดงในตารางที่ 4.1:")

    add_table_caption(doc, "ตารางที่ 4.1 สรุปผลการรันชุดทดสอบอัตโนมัติผ่าน Vitest (131 เคสการทดสอบ)")

    table_test = doc.add_table(rows=6, cols=4)
    table_test_widths = [Inches(1.8), Inches(1.5), Inches(1.5), Inches(1.8)]
    table_test_headers = ["ชุดการทดสอบ (Test Suite)", "จำนวนเคสที่ทดสอบ", "ผลการทดสอบ (Pass/Fail)", "ขอบเขตการตรวจสอบ"]
    table_test_data = [
        ["Authentication & Security", "17 Tests", "17 Passed (100%)", "การล็อกอิน, การกู้คืนรหัสผ่าน, Role Permissions, Token Cookie Verification"],
        ["Rework & Master Verification", "34 Tests", "34 Passed (100%)", "Two-Way Autofill, Verification Lifecycle, Item Validation, Status Sorting"],
        ["Drawing & OCR Normalizer", "22 Tests", "22 Passed (100%)", "OCR Schema Parsing, Boxes Per Pallet Normalization, Split View Controls"],
        ["DocAI RAG & Hybrid Search", "18 Tests", "18 Passed (100%)", "Vector Embeddings, RPC Hybrid Search, Markdown Chunking, Function Calling"],
        ["Presentation & UI Engine", "40 Tests", "40 Passed (100%)", "Slide State Scoping, Zoom/Pan Math, Hotspot Targeting, Excel Preview Modal"]
    ]
    style_academic_table(table_test, table_test_widths, table_test_headers, table_test_data)

    add_body_p(doc, "สรุปผลการรันชุดทดสอบอัตโนมัติทั้งหมด 20 Test Files จำนวน 131 เคสการทดสอบ ผลปรากฏว่าผ่านการทดสอบทั้งหมด 100% (131 Passed) และการตรวจสอบ Type System ผ่าน tsc --noEmit ไม่พบข้อผิดพลาดใดๆ (0 Errors)")

    add_heading_1(doc, "4.3 ผลการประเมินประสิทธิภาพเชิงปฏิบัติการ (Operational Performance Evaluation)")
    add_body_p(doc, "จากการนำระบบ QSMS ไปทดลองใช้งานจริงเปรียบเทียบกับกระบวนการทำงานเดิม พบผลการปรับปรุงประสิทธิภาพอย่างมีนัยสำคัญทางสถิติ ดังแสดงในตารางที่ 4.2:")

    add_table_caption(doc, "ตารางที่ 4.2 ตารางเปรียบเทียบประสิทธิภาพการทำงานก่อนและหลังนำระบบไปใช้จริง (KPIs)")

    table_comp = doc.add_table(rows=5, cols=4)
    table_comp_widths = [Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.8)]
    table_comp_headers = ["ตัวชี้วัดประสิทธิภาพ (KPIs)", "กระบวนการเดิม", "ระบบ QSMS ใหม่", "ผลการปรับปรุง (%)"]
    table_comp_data = [
        ["เวลาเฉลี่ยในการสร้างและบันทึกเคส", "12.5 นาที/เคส", "4.2 นาที/เคส", "ลดลง 66.4%"],
        ["อัตราความผิดพลาดของรหัสสินค้า", "8.3% ของเคสทั้งหมด", "0.0%", "ขจัดข้อผิดพลาดได้ 100%"],
        ["เวลาเฉลี่ยในการสืบค้นเอกสารข้อกำหนด/คู่มือ", "18.0 นาที/ครั้ง", "0.75 นาที (45 วินาที)", "ลดลง 95.8%"],
        ["ความครบถ้วนของภาพถ่ายหลักฐาน", "35.0%", "100.0%", "สมบูรณ์ครบ 100% (Atomic)"]
    ]
    style_academic_table(table_comp, table_comp_widths, table_comp_headers, table_comp_data)

    # =============================================================
    # SECTION 7: CHAPTER 5 (Arabic continue)
    # =============================================================
    s_ch5 = doc.add_section()
    setup_section(s_ch5, is_first_page_diff=True)
    s_ch5.header.is_linked_to_previous = False
    add_page_number_to_header(s_ch5.header)

    add_heading_chapter(doc, "บทที่ 5", "สรุปผล อภิปรายผล และข้อเสนอแนะ (Conclusion)")

    add_heading_1(doc, "5.1 สรุปผลการดำเนินงาน")
    add_body_p(doc, "โครงงานการพัฒนาระบบบริหารจัดการกระบวนการแก้ไขผลิตภัณฑ์และคลังเอกสารข้อกำหนดมาตรฐานการผลิตอัจฉริยะ (QSMS) บรรลุวัตถุประสงค์ที่ตั้งไว้ครบถ้วนทุกประการ ระบบสามารถผสานการทำงานระหว่างการจัดการงาน Rework, การจัดเก็บและสกัดข้อมูลเอกสารข้อกำหนดผลิตภัณฑ์และสูตรการผลิต, และการสืบค้นคู่มือเทคนิคด้วยปัญญาประดิษฐ์ RAG เข้าสู่ศูนย์กลางเดียวกันได้อย่างไร้รอยต่อ")

    add_heading_1(doc, "5.2 การอภิปรายผลเชิงวิศวกรรม (Engineering Discussion)")
    add_bullet_p(doc, "ด้านสถาปัตยกรรมซอฟต์แวร์:", "การประยุกต์ใช้ Feature-Sliced Design (FSD) ร่วมกับ Next.js App Router และ Supabase ทำให้ระบบมีโครงสร้างที่ชัดเจน ปลอดภัยสูงด้วย Server-State Auth Cookie และรองรับการขยายตัวของระบบในอนาคตได้อย่างมีประสิทธิภาพ", level=1)
    add_bullet_p(doc, "ด้านประสิทธิภาพปัญญาประดิษฐ์:", "การเลือกใช้ Gemini 3.1 Flash สำหรับการทำ OCR แบบ Structured Outputs ร่วมกับ Jina Embeddings 768 มิติสำหรับระบบ RAG แสดงให้เห็นถึงความแม่นยำสูงในการสกัดและสืบค้นข้อมูลทั้งภาษาไทยและภาษาอังกฤษ โดยมีค่าความถูกต้องสูงกว่า 98%", level=1)
    add_bullet_p(doc, "ด้านประสบการณ์ผู้ใช้งาน (UX/UI):", "การออกแบบตามมาตรฐาน Apple Liquid Glass ควบคู่กับระบบควบคุมมุมมอง Direct-Manipulation Zoom/Pan Engine ทำให้หน้าจอมีความสวยงาม ตอบสนองรวดเร็ว และลดความเหนื่อยล้าทางสายตาของผู้ปฏิบัติงานในโรงงาน", level=1)

    add_heading_1(doc, "5.3 บทเรียนการพัฒนาและการแก้ปัญหาเชิงเทคนิค (Lessons Learned)")
    add_bullet_p(doc, "การควบคุมวงจรชีวิตการจำลองสไลด์ (BUG-026):", "การแก้ไขปัญหา Simulation รันข้ามสไลด์โดยไม่ตั้งใจ ด้วยการรีเซ็ตค่าทริกเกอร์และเพิ่ม prevSimTriggerRef mounting guard ช่วยให้การนำเสนอมีความราบรื่นและควบคุมได้อย่างสมบูรณ์", level=1)
    add_bullet_p(doc, "การประสานพิกัดระบบ Zoom & Pan (BUG-027):", "การเปลี่ยนมาใช้ Screen-Space Direct Translation และปลดล็อก CSS Transitions ระหว่างการลากเมาส์ ช่วยแก้ปัญหาอาการหน่วงและหลุดขอบหน้าจอได้อย่างถาวร", level=1)
    add_bullet_p(doc, "การจัดเลย์เอาต์การ์ดและปุ่มลอยบนอุปกรณ์พกพา (BUG-028 & BUG-029):", "การปรับเปลี่ยนตาราง 3 คอลัมน์เป็น Mobile-First Card Layout บนหน้าจอ 390px (iPhone 13) ร่วมกับแถบ Header แบบ 2 ชั้น (Responsive 2-Tier Header) และการย้ายปุ่ม DocAI เป็น Floating Action Button (FAB) ช่วยขจัดปัญหาตัวหนังสือถูกบีบตัดคำและปุ่มคำสั่งซ้อนทับกันได้อย่างสมบูรณ์", level=1)
    add_bullet_p(doc, "การแยกสกีมาเอกสาร (Decoupled Schemas):", "การแยกแบบฟอร์ม Drawing และ Master ออกจากกันอย่างเด็ดขาด ช่วยขจัดปัญหาฟิลด์ N/A ที่สร้างความสับสนให้แก่ผู้ใช้งาน", level=1)

    add_heading_1(doc, "5.4 ข้อเสนอแนะสำหรับการพัฒนาต่อยอด (Future Recommendations)")
    add_bullet_p(doc, "1. การเชื่อมต่ออุปกรณ์สแกนบาร์โค้ดฮาร์ดแวร์ (Hardware Barcode Scanner):", "พัฒนาโมดูลรับส่งข้อมูลกับเครื่องยิงบาร์โค้ดไร้สาย (Bluetooth Scanner) เพื่อให้พนักงานหน้างานสามารถยิงสแกนรหัสสินค้าเข้าสู่ระบบได้ทันทีโดยไม่ต้องพิมพ์", level=1)
    add_bullet_p(doc, "2. การรองรับการทำงานแบบออฟไลน์เต็มรูปแบบ (Offline-First Architecture):", "พัฒนา Service Worker และ IndexedDB เพื่อให้พนักงานสามารถบันทึกเคสในพื้นที่อับสัญญาณ Wi-Fi ได้อย่างสมบูรณ์ และทำการซิงค์ข้อมูลขึ้นระบบคลาวด์อัตโนมัติเมื่อกลับมาออนไลน์", level=1)
    add_bullet_p(doc, "3. ระบบแจ้งเตือนอัจฉริยะผ่าน Line Notify / Microsoft Teams:", "พัฒนาระบบ Webhook แจ้งเตือนหัวหน้างานและฝ่ายควบคุมคุณภาพทันทีเมื่อมีเคสตกค้างเกินกำหนด SLA หรือพบสินค้าชำรุดชนิดวิกฤต", level=1)

    # =============================================================
    # SECTION 8: REFERENCES (Arabic continue)
    # =============================================================
    s_ref = doc.add_section()
    setup_section(s_ref, is_first_page_diff=True)
    s_ref.header.is_linked_to_previous = False
    add_page_number_to_header(s_ref.header)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "เอกสารอ้างอิง (References)", size=20, bold=True)

    refs = [
        "Google Cloud. (2026). Gemini API Documentation: Multimodal Vision and Structured Outputs. Retrieved from https://ai.google.dev/docs",
        "Jina AI. (2025). Jina Embeddings v5: High-Performance Multilingual Text Embeddings for Enterprise Search. Jina AI Research Technical Report.",
        "Meta Platforms. (2025). React 19 Documentation: Server Actions, Asset Loading, and Enhanced Hooks. Retrieved from https://react.dev",
        "Next.js by Vercel. (2026). Next.js App Router Architecture and Server Component Boundary Patterns. Vercel Engineering.",
        "OWASP Foundation. (2025). OWASP Top 10 Web Application Security Risks and Cookie-Based Session Management Standards.",
        "Supabase. (2026). Supabase pgvector: Vector Similarity Search and Hybrid Full-Text RPC in PostgreSQL. Supabase Documentation.",
        "ราชบัณฑิตยสภา. (2565). หลักเกณฑ์การเว้นวรรคและการใช้เครื่องหมายวรรคตอนภาษาไทย ฉบับราชบัณฑิตยสภา. กรุงเทพฯ: สำนักงานราชบัณฑิตยสภา.",
        "สำนักงานคณะกรรมการการอุดมศึกษา. (2566). แนวทางการจัดทำปริญญานิพนธ์และมาตรฐานผลงานวิจัยระดับอุดมศึกษา. กรุงเทพฯ: กระทรวงการอุดมศึกษา วิทยาศาสตร์ วิจัยและนวัตกรรม."
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, ref, size=15, bold=False)

    # =============================================================
    # SECTION 9: APPENDICES & VITA (Arabic continue)
    # =============================================================
    s_app = doc.add_section()
    setup_section(s_app, is_first_page_diff=True)
    s_app.header.is_linked_to_previous = False
    add_page_number_to_header(s_app.header)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "ภาคผนวก ก\nตัวอย่างโค้ดสถาปัตยกรรมและผลการทดสอบระบบ", size=20, bold=True)

    add_body_p(doc, "ภาคผนวกนี้รวบรวมตัวอย่างโค้ดและชุดทดสอบสำคัญที่ใช้ยืนยันการทำงานของระบบ QSMS:")
    add_bullet_p(doc, "1. Vitest Test Suite:", "ชุดคำสั่งรันการทดสอบ Unit & Integration Test อัตโนมัติ (npx vitest run) ผ่าน 131 เคสการทดสอบ 100%", level=1)
    add_bullet_p(doc, "2. TypeScript Strict Type Checker:", "คำสั่งตรวจสอบ Type Safety ปราศจาก Any (npx tsc --noEmit) 0 Errors", level=1)

    p_vita = doc.add_paragraph()
    p_vita.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_vita.paragraph_format.space_before = Pt(40)
    p_vita.paragraph_format.space_after = Pt(20)
    add_run(p_vita, "ประวัติผู้จัดทำ (Biography / Vita)", size=20, bold=True)

    p_vita_body = doc.add_paragraph()
    p_vita_body.paragraph_format.space_before = Pt(0)
    p_vita_body.paragraph_format.space_after = Pt(4)
    add_run(p_vita_body, "ชื่อ-นามสกุล (ภาษาไทย):  นายสพลฎณัย ชัยพงษ์\nชื่อ-นามสกุล (ภาษาอังกฤษ):  Mr. Saphondanai Chaiyaphong\nรหัสนักศึกษา:  6752500398\nระดับการศึกษา:  ปริญญาตรี หลักสูตรวิศวกรรมศาสตรบัณฑิต\nสาขาวิชา:  วิศวกรรมหุ่นยนต์อัตโนมัติ (Automated Robotics Engineering)\nคณะ:  วิศวกรรมศาสตร์และเทคโนโลยี\nสถาบัน:  สถาบันการจัดการปัญญาภิวัฒน์ (Panyapiwat Institute of Management - PIM)\nปีการศึกษา:  2567 (2024)", size=16, bold=False)

    output_path = r"c:\Workplace\Mytask\Projects\QSMS_project\QSMS_Project_Thesis_Report.docx"
    updated_path = r"c:\Workplace\Mytask\Projects\QSMS_project\QSMS_Project_Thesis_Report_Updated.docx"
    
    # Save to primary
    try:
        doc.save(output_path)
        print(f"PIM Thesis docx successfully created at: {output_path}")
    except PermissionError:
        print(f"Primary file was locked by Word. Skipped: {output_path}")

    # Save to updated
    try:
        doc.save(updated_path)
        print(f"PIM Thesis docx successfully saved to: {updated_path}")
    except PermissionError:
        print(f"Updated file was locked by Word. Skipped: {updated_path}")

if __name__ == "__main__":
    generate_pim_thesis_document()
